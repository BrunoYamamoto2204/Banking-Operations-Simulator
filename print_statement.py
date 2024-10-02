from datetime import datetime
from reportlab.pdfgen import canvas
from docx import Document


def statement(user,ini_bal,deposit,total_deposits,withdrawal,total_withdrawals,initial_balance): #Criar pdf
    doc = Document()
    agora = datetime.now()
    data_hora_formatada = agora.strftime("%Y-%m-%d %H:%M:%S")

    doc.add_paragraph(f'<{data_hora_formatada}>')
    doc.add_paragraph(f'STATE BANK')
    doc.add_paragraph(f'Here is your Statement, {user}')
    doc.add_paragraph(f'Initial Balance: ${ini_bal}')
    doc.add_paragraph(f'Deposits Made: {deposit}')
    doc.add_paragraph(f'Total Deposits: ${total_deposits}')
    doc.add_paragraph(f'Withdrawals Made: {withdrawal}')
    doc.add_paragraph(f'Total Withdrawals: ${total_withdrawals}')
    doc.add_paragraph(f'Ending Balance: ${initial_balance}')

    doc.add_paragraph(f'STATE BANK')
    doc.add_paragraph(f'123 MAIN STREET, LOS ANGELES, CA')
    doc.save('statement.docx')
    convert_docx_to_pdf("statement.docx", "statement.pdf")
    print('-' * 40)

def convert_docx_to_pdf(docx_path, pdf_path):
    c = canvas.Canvas(pdf_path)

    document = Document(docx_path)
    y = 800

    for paragraph in document.paragraphs:
        c.drawString(100, y, paragraph.text)
        y -= 20

    c.save()